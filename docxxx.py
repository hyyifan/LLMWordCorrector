import os
import json
import re
from docx import Document
import win32com.client as win32
from openai import OpenAI
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
import pythoncom
import threading
import tempfile

class DocxReview:
    def __init__(self, doc_path, key, split_length=1000, max_retries=3, retry_delay=5, max_workers=5):
        self.client = OpenAI(
            api_key=key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        )
        self.doc_path = os.path.abspath(doc_path)
        self.split_length = split_length
        self.find_count = 0
        self.notfind_count = 0
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.max_workers = max_workers
        self.progress_lock = Lock()
        self.total_chunks = 0
        self.processed_chunks = 0
        self.last_printed_progress = -1  # 新增：记录上次打印的进度
        self.blocked_words = ["共产党", "习近平", "毛泽东", "中共","反革命","六四","天安门事件","法轮功","台独","西藏独立","新疆独立","民主革命","政治改革"]

    def read_document(self):
        file_extension = os.path.splitext(self.doc_path)[1].lower()
        
        if file_extension == '.docx':
            return self._read_docx()
        elif file_extension == '.doc':
            if self._is_word_installed():
                print("【注意】正在将 .doc 文件转换为 .docx 格式以进行读取...")
                temp_docx_path = self._convert_doc_to_docx(self.doc_path)
                try:
                    result = self._read_docx(temp_docx_path)
                    print("【成功】已成功读取转换后的文件内容。")
                    return result
                finally:
                    # 删除临时文件
                    if os.path.exists(temp_docx_path):
                        os.remove(temp_docx_path)
                        print("【清理】临时转换文件已删除。")
            else:
                error_msg = "【错误】无法读取 .doc 文件。您的电脑未安装 Microsoft Word 软件。请安装 Word 或使用 .docx 格式的文件。"
                print(error_msg)
        else:
            error_msg = "【错误】不支持的文件格式。请使用 .doc 或 .docx 格式的文件。"
            print(error_msg)

    def _is_word_installed(self):
        try:
            win32.Dispatch("Word.Application")
            return True
        except pythoncom.com_error:
            return False

    def _read_docx(self, docx_path=None):
        doc = Document(docx_path or self.doc_path)
        content = []
        last_content = None  # 用于存储上一个添加的内容
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cleaned_text = cell.text.strip()
                        if cleaned_text != last_content:
                            content.append(cleaned_text)
                            last_content = cleaned_text
        
        # 读取段落内容
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned_text = para.text.strip()
                if cleaned_text != last_content:
                    content.append(cleaned_text)
                    last_content = cleaned_text
        
        return "\n".join(content)

    def _convert_doc_to_docx(self, doc_path):
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(doc_path)
            # 创建临时文件路径
            temp_dir = tempfile.gettempdir()
            docx_path = os.path.join(temp_dir, "temp_converted.docx")
            # 另存为 .docx 格式
            doc.SaveAs(docx_path, FileFormat=16)  # FileFormat=16 表示 .docx 格式
            return docx_path
        finally:
            doc.Close()
            word.Quit()

    # def _read_doc(self):
    #     word = win32.Dispatch("Word.Application")
    #     word.Visible = False
    #     try:
    #         doc = word.Documents.Open(self.doc_path)
    #         content = doc.Range().Text
    #         # 移除多余的空白行和空格
    #         cleaned_content = "\n".join(line.strip() for line in content.split("\n") if line.strip())
    #         return cleaned_content
    #     finally:
    #         doc.Close()
    #         word.Quit()

    def remove_blocked_words(self, text):
        for word in self.blocked_words:
            text = text.replace(word, "")
        return text
    

    def split_text(self, full_word):
        split_pattern = r'(\n|。|！|？|；)'
        segments = re.split(split_pattern, full_word)

        processed_segments = []
        for i in range(0, len(segments) - 1, 2):
            segment = segments[i] + segments[i + 1]
            processed_segments.append(segment)
        if len(segments) % 2 != 0:
            processed_segments.append(segments[-1])  # 处理剩余部分

        # 按字符长度合并
        final_segments = []
        current_segment = ""
        for segment in processed_segments:
            if len(current_segment) + len(segment) <= self.split_length:
                current_segment += segment
            else:
                final_segments.append(current_segment)
                current_segment = segment
        if current_segment:
            final_segments.append(current_segment)

        return final_segments
    
    def extract_and_parse_json(self,text):
        # 提取 {} 之间的内容
        matches = re.findall(r'\{[^{}]*\}', text)
        
        if not matches:
            print("警告: 在输入字符串中没有找到有效的JSON。")
            return []
        
        # 尝试解析每个匹配项
        parsed_jsons = []
        for i, match in enumerate(matches):
            try:
                parsed_json = json.loads(match)
                parsed_jsons.append(parsed_json)
            except json.JSONDecodeError as e:
                print(f"警告: 第 {i+1} 个JSON对象解析失败: {str(e)}")
                continue
        
        if not parsed_jsons:
            print("警告: 没有成功解析任何JSON对象。")
            return []
        
        # 筛选"是否错误"为"是"的内容
        errors = [item for item in parsed_jsons if item.get("是否错误") == "是"]
        
        if not errors:
            print("【信息】: 在解析的JSON对象中没有找到“错误段落”。")
        
        return errors

    def word_review(self, json_data):
        try:
            # 启动 Word 应用程序
            word_app = win32.Dispatch("Word.Application")
            word_app.Visible = False  # 设置为True可以看到Word窗口

            # 打开文档
            doc = word_app.Documents.Open(self.doc_path)

            # 获取文档中的选择对象和查找功能
            selection = word_app.Selection
            find = selection.Find

            # 用于存储未能批注的内容
            manual_review_items = []

            # 遍历 json_data 并查找指定文本
            for index, item in enumerate(json_data, 1):
                error_position = item.get("错误位置")
                error_reason = item.get("错误原因")
                suggestion = item.get("修改意见")
                
                if not all([error_position, error_reason, suggestion]):
                    print(f"警告: JSON 数据不完整: {item}")
                    continue

                selection.HomeKey(Unit=6)  # 将光标重置到文档开头
                find.Text = error_position  # 设置要查找的文本
                find.Execute()  # 执行查找操作

                # 如果找到匹配的文本，尝试添加评论
                if find.Found:
                    try:
                        comment_text = f"【错误原因】{error_reason}。【修改意见】{suggestion}"
                        doc.Comments.Add(selection.Range, comment_text)
                        self.find_count += 1
                    except Exception as e:
                        # 捕获批注失败的错误并添加到手动审阅列表
                        manual_review_items.append({
                            "index": index,
                            "error_position": error_position,
                            "error_reason": error_reason,
                            "suggestion": suggestion,
                            "reason": str(e)
                        })
                        self.notfind_count += 1
                else:
                    # 未找到文本时，添加到手动审阅列表
                    manual_review_items.append({
                        "index": index,
                        "error_position": error_position,
                        "error_reason": error_reason,
                        "suggestion": suggestion,
                        "reason": "未找到指定文本"
                    })
                    self.notfind_count += 1

            # 保存文档
            doc.Save()

            # 输出需要手动审阅的内容
            if manual_review_items:
                print("\n【需要手动审阅的内容】")
                for item in manual_review_items:
                    print(f"\n{item['index']}. 【错误位置】: {item['error_position']}")
                    print(f"   【错误原因】: {item['error_reason']}")
                    print(f"   【修改意见】: {item['suggestion']}")
                    print(f"   【未能自动批注原因】: {item['reason']}")

        except Exception as e:
            print(f"发生错误: {e}")
        finally:
            # 确保文档关闭和 Word 应用退出，即使在出现异常时也会执行
            if 'doc' in locals():
                doc.Close()  # 关闭文档
            word_app.Quit()  # 退出应用程序

    def chat_single(self, item):
        for attempt in range(self.max_retries):
            try:
                query = item
                completion = self.client.chat.completions.create(
                model= "qwen-plus-0919",
                ###"qwen-max-0919",
                ###"qwen-plus-0919"
                messages=[
                    {
                        'role': 'system',
                        'content': """
                        <promopts>
                        - Role: 错词错句校验专家
                        - Background: 用户需要审核提交的文字内容，专注于纠正错别字和用词不当。
                        - Profile: 你是一位资深的错词错句校验专家，具备敏锐的语言洞察力和丰富的校对经验，能够迅速识别并纠正文本中的错误。
                        - Skills: 你擅长识别错别字、同音异义词误用，以及在上下文中准确判断词汇的恰当性，确保文本的准确性和专业性。
                        - Goals: 确保文本中所有错别字和用词不当得到纠正，同时保持原文意义和结构的完整性。
                        - Constrains: 工作范畴不涉及句法调整、语义澄清或结构优化；专注于输出存在错误的句子和错误原因。
                        - OutputFormat: 列出原句和修正建议，确保原句完整不删减，便于搜索。
                        - Workflow:
                        1. 依据句子分隔符号(。！？；等)和分段符号"\n",将文本分割为独立的句子。
                        2. 逐句审查，识别错别字或用词不当。
                        3. 在上下文中判断词汇的恰当性。
                        4. 对存在错误的词语、句子,请列出原句，并标记存在错误，提供修改意见。
                        5. 对不存在错误的词语，请你直接忽略该句子，开始审阅下一个句子。
                        6. 确保输出中不夹杂任何XML标签。
                        7. 确保输出样式不添加任意符号（包括#*-）
                        8. 确保输出中不夹在任何本promopt的内容。
                        9. 请严格根据examples的json格式进行输出。
                        - Examples:
                        - 例子1：
                            {"原句"："党的二十大报告强调，要提高防灾减灾救灾和重大突发公共事件处置保障能力，加强国家区域应激力量建设！"，
                            "是否错误"："是"，
                            "错误原因"："存在错别字"，
                            "错误位置"："加强国家区域应激力量建设"，
                            "修改意见"："应修正为 “加强国家区域应急力量建设”"}
                            
                        - 例子2：
                            {"原句"："本项目泥在平潭海坛湾外海布设浮标，根据《国土空间调查、规划、用途管制用地用海分类指南》，本项目用海类型为“特殊用海”中的“科研教育用海”；"，
                            "是否错误"："是"，
                            "错误原因"："存在错别字"，
                            "错误位置"："本项目泥在平潭海坛湾外海布设浮标"，
                            "修改意见"："应修正为 "本项目拟在平潭海坛湾外海布设浮标""}
                            
                        - 例子3：
                            {"原句"："赤潮是平潭综合实验区主要的海洋生态灾害类型之一，在2011-2015年，全省工发现赤潮事件20余起，其中平潭综合实验区每年萍君发生3起赤潮事件，发生时间段为4月至6月。"，
                            "是否错误"："是"，
                            "错误原因"："存在错别字，错误用词"，
                            "错误位置"："全省工发现赤潮事件20余起，其中平潭综合实验区每年萍君发生3起赤潮事件"，
                            "修改意见"：""全省工发现赤潮事件20余起" 应修正为 "全省共发现赤潮事件20余起"；“平潭综合实验区每年萍君发生3起赤潮事件”应修正为“其中平潭综合实验区每年平均发生3起赤潮事件”"}              

                        - 例子4：
                            {"原句"："除此之外，为进一步加强城市智能化程度，福州市还额外开展了xx项重点项目，建设基础设施xx个，投入资金规模达xx亿元。"，
                            "是否错误"："是"，
                            "错误原因"："词语搭配不当"，
                            "错误位置"："为进一步加强城市智能化程度"，
                            "修改意见"："应修正为“为进一步提升城市智能化水平”"}

                        - 例子5：
                            {"原句"："试点建设期间，项目获得众多奖项xxx。"，
                            "是否错误"："是 "，
                            "错误原因"："存在无意义符号" ，
                            "错误位置"："项目获得众多奖项xxx"，
                            "修改意见"："应修正为“项目获得众多奖项”"}

                        - 例子6：
                            {"原句"："试点建设期间，福州市以基础设施信息化、数字化为基础，打造了依稀列前端基础设施和智能化处理系统，极大地丰富了市民的智能生活体验。"，
                            "是否错误"："是"，
                            "错误原因"："用词不当"，
                            "错误位置"："打造了一系列前端基础设施和智能化处理系统"，
                            "修改意见"："应修正为“打造了一系列前端基础设施和智能化处理系统”"}
                                                    
                        - 例子7：
                            {"原句"："党的二十大报告强调，要提高防灾减灾救灾和重大突发公共事件处置保障能力，加强国家区域应急力量建设。近年来，各级政府部门对海洋防灾减灾工作的重视程度明显提升。"，
                            "是否错误"："无"，
                            "错误原因"："无"，
                            "错误位置"："无"，
                            "修改意见"："无"}
                                    
                        - 例子8：
                            {"原句"："项目概况"，
                            "是否错误"："无"，
                            "错误原因"："无"，
                            "错误位置"："无"，
                            "修改意见"："无"}
                            
                        - 例子9：
                            {"原句"："风险防范对策措施"，
                            "是否错误"："无"，
                            "错误原因"："无"，
                            "错误位置"："无"，
                            "修改意见"："无"}
                            </promopts>
                        """
                    },
                    {
                        'role': 'user',
                        'content': f'请分析以下内容是否有错字、错词、错句：{query}'
                    }
                ],
                )
                return completion.model_dump_json()
            except Exception as e:
                if attempt < self.max_retries - 1:
                    print(f"尝试 {attempt + 1} 失败，等待 {self.retry_delay} 秒后重试...")
                    time.sleep(self.retry_delay)
                else:
                    print(f"【错误】{e}")

    def update_progress(self):
        with self.progress_lock:
            self.processed_chunks += 1
            progress = (self.processed_chunks / self.total_chunks) * 100
            # 每增加5%打印一次进度
            if int(progress) // 5 > self.last_printed_progress // 5:
                print(f"【进度】完成 {int(progress)}%")
                self.last_printed_progress = int(progress)

    def process_chunk(self, chunk):
        thread_id = threading.get_ident()
        # print(f"线程 {thread_id} 开始处理文本块")
        try:
            # start_time = time.time()
            llm_result = json.loads(self.chat_single(chunk))['choices'][0]['message']['content']
            # end_time = time.time()
            # processing_time = end_time - start_time
            # print(f"线程 {thread_id} 完成处理，耗时 {processing_time:.2f} 秒")
            self.update_progress()
            return llm_result
        except Exception as e:
            # print(f"线程 {thread_id} 处理文本块时发生错误: {str(e)}")
            self.update_progress()
            return f"处理文本块时发生错误: {str(e)}"

    def run(self):
        print("【开始运行】")
        print(f"当前参数设置：")
        print(f"- 单次分析字数: {self.split_length}")
        print(f"- 最大重试次数: {self.max_retries}")
        print(f"- 重试延迟: {self.retry_delay}秒")
        print(f"- 工作线程数: {self.max_workers}")
        print(f"- 已启用屏蔽词过滤，共 {len(self.blocked_words)} 个屏蔽词")
        print("------------------------")

        doc_content = self.remove_blocked_words(self.read_document())
        print("已读取了文档内容")
        word_split = self.split_text(doc_content)
        self.total_chunks = len(word_split)
        chat_all_result = ''
        errors = []

        print("【大模型开始分析】")
        print(f"文本共包含{len(doc_content)}个文字，将进行{self.total_chunks}段分析")
        print("请稍等...")
        print(f"【进度】完成 {int(0)}%")

        start_time = time.time()
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_chunk = {executor.submit(self.process_chunk, chunk): i for i, chunk in enumerate(word_split)}
            
            for future in as_completed(future_to_chunk):
                chunk_index = future_to_chunk[future]
                try:
                    result = future.result()
                    if result.startswith("处理文本块时发生错误"):
                        errors.append(f"在处理第{chunk_index+1}段文字时发生错误: {result}")
                    else:
                        chat_all_result += result + "\n\n"
                except Exception as exc:
                    errors.append(f"在处理第{chunk_index+1}段文字时发生错误: {str(exc)}")

        end_time = time.time()
        total_time = end_time - start_time
        print(f"分析完成，总耗时 {total_time:.2f} 秒")

        # 将文本转换为json
        llm_json = self.extract_and_parse_json(chat_all_result) 
        print(f"本文本共存在{len(llm_json)}个错误\n")

        print("【批注】")
        print("正在给文档增加批注...\n")
        self.word_review(llm_json)

        print("【批注结束！】\n")

        print("【批注结果】")
        print(f"已批注的文本数量: {self.find_count}\n")
        print(f"未找到位置，因而没有批注的数量: {self.notfind_count}\n")
        
        if errors:
            print("【处理过程中遇到的问题】")
            for error in errors:
                print(error)
        
        print("结束，感谢使用!")

# 使用示例
if __name__ == "__main__":
    doc_path = r'D:\test\aaaa1.docx'
    api_key = "your_api_key_here"  # 请替换为您的实际API密钥
    docx_reviewer = DocxReview(doc_path, api_key, max_workers=5)
    docx_reviewer.run()